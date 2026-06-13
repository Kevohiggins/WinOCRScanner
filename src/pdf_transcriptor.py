import wx
import os
import threading
import time
import fitz  # PyMuPDF
import numpy as np
import cv2
from docx import Document
import pandas as pd

from ocr_engine import OCREngine
from table_engine import TableEngine
from layout_engine import LayoutEngine
from config import load_config
from translator import translator_instance

class TranscriptorFrame(wx.Frame):
    def __init__(self, config=None):
        super().__init__(None, title="Modo Documentos - PaddleOCR Scanner", size=(700, 750))
        self.config = config or load_config()
        self.ocr = None
        self.table = None
        self.layout = None
        self.InitUI()

    def InitUI(self):
        self.main_panel = wx.Panel(self)
        self.main_sizer = wx.BoxSizer(wx.VERTICAL)

        # Header
        self.header_panel = wx.Panel(self.main_panel)
        self.header_panel.SetBackgroundColour(wx.Colour(45, 45, 45))
        header_sizer = wx.BoxSizer(wx.VERTICAL)
        title = wx.StaticText(self.header_panel, label="TRANSCRIPTOR/TRADUCTOR DE DOCUMENTOS")
        title.SetForegroundColour(wx.WHITE)
        font = title.GetFont()
        font.SetPointSize(14)
        font.SetWeight(wx.FONTWEIGHT_BOLD)
        title.SetFont(font)
        header_sizer.Add(title, 0, wx.ALL | wx.CENTER, 15)
        self.header_panel.SetSizer(header_sizer)
        self.main_sizer.Add(self.header_panel, 0, wx.EXPAND)

        content_sizer = wx.BoxSizer(wx.VERTICAL)

        # 1. File Selection
        file_box = wx.StaticBox(self.main_panel, label="Archivo de Origen (Soporta PDF, EPUB, MOBI, XPS, CBZ, FB2, etc.)")
        file_sizer = wx.StaticBoxSizer(file_box, wx.HORIZONTAL)
        self.txt_path = wx.TextCtrl(self.main_panel, style=wx.TE_READONLY)
        btn_browse = wx.Button(self.main_panel, label="Buscar Archivo")
        btn_browse.Bind(wx.EVT_BUTTON, self.OnBrowse)
        file_sizer.Add(self.txt_path, 1, wx.EXPAND | wx.ALL, 5)
        file_sizer.Add(btn_browse, 0, wx.ALL, 5)
        content_sizer.Add(file_sizer, 0, wx.EXPAND | wx.ALL, 10)

        # 2. Modes and Options
        opt_box = wx.StaticBox(self.main_panel, label="Opciones de Procesamiento")
        opt_sizer = wx.StaticBoxSizer(opt_box, wx.VERTICAL)
        
        # Operation mode
        mode_sizer = wx.BoxSizer(wx.HORIZONTAL)
        mode_sizer.Add(wx.StaticText(self.main_panel, label="Modo de Operación:"), 0, wx.ALIGN_CENTER_VERTICAL | wx.ALL, 5)
        self.cb_mode = wx.Choice(self.main_panel, choices=["Solo Transcripción", "Solo Traducción", "Híbrido (Original + Traducción)"])
        self.cb_mode.SetSelection(2) # Default to Hybrid
        mode_sizer.Add(self.cb_mode, 1, wx.EXPAND | wx.ALL, 5)
        opt_sizer.Add(mode_sizer, 0, wx.EXPAND)
        
        # Checkboxes
        self.chk_tables = wx.CheckBox(self.main_panel, label="Forzar detección OCR de tablas en imágenes (Si están escaneadas)")
        self.chk_tables.SetValue(True)
        self.chk_page_breaks = wx.CheckBox(self.main_panel, label="Mantener saltos de página")
        self.chk_page_breaks.SetValue(True)
        opt_sizer.Add(self.chk_tables, 0, wx.ALL, 5)
        opt_sizer.Add(self.chk_page_breaks, 0, wx.ALL, 5)
        content_sizer.Add(opt_sizer, 0, wx.EXPAND | wx.ALL, 10)

        # 3. Output Format
        out_box = wx.StaticBox(self.main_panel, label="Formato de Salida")
        out_sizer = wx.StaticBoxSizer(out_box, wx.HORIZONTAL)
        self.rb_txt = wx.RadioButton(self.main_panel, label="Texto (.txt)", style=wx.RB_GROUP)
        self.rb_word = wx.RadioButton(self.main_panel, label="Word (.docx)")
        self.rb_excel = wx.RadioButton(self.main_panel, label="Excel (.xlsx)")
        self.rb_word.SetValue(True)
        out_sizer.Add(self.rb_txt, 1, wx.ALL, 5)
        out_sizer.Add(self.rb_word, 1, wx.ALL, 5)
        out_sizer.Add(self.rb_excel, 1, wx.ALL, 5)
        content_sizer.Add(out_sizer, 0, wx.EXPAND | wx.ALL, 10)

        # 4. Progress and Logs
        self.gauge = wx.Gauge(self.main_panel, range=100)
        self.log_ctrl = wx.TextCtrl(self.main_panel, style=wx.TE_MULTILINE | wx.TE_READONLY | wx.HSCROLL)
        self.log_ctrl.SetBackgroundColour(wx.Colour(245, 245, 245))
        content_sizer.Add(self.gauge, 0, wx.EXPAND | wx.LEFT | wx.RIGHT | wx.TOP, 15)
        content_sizer.Add(self.log_ctrl, 1, wx.EXPAND | wx.ALL, 10)

        # 5. Run Button
        self.btn_run = wx.Button(self.main_panel, label="INICIAR PROCESAMIENTO", size=(-1, 60))
        self.btn_run.SetBackgroundColour(wx.Colour(0, 120, 215))
        self.btn_run.SetForegroundColour(wx.WHITE)
        font_btn = self.btn_run.GetFont()
        font_btn.SetWeight(wx.FONTWEIGHT_BOLD)
        self.btn_run.SetFont(font_btn)
        self.btn_run.Bind(wx.EVT_BUTTON, self.OnRun)
        content_sizer.Add(self.btn_run, 0, wx.EXPAND | wx.ALL, 15)

        self.main_sizer.Add(content_sizer, 1, wx.EXPAND)
        self.main_panel.SetSizer(self.main_sizer)
        self.Centre()

    def Log(self, message):
        wx.CallAfter(self.log_ctrl.AppendText, f"[{time.strftime('%H:%M:%S')}] {message}\n")

    def OnBrowse(self, event):
        wildcard = "Archivos Soportados (*.pdf;*.epub;*.mobi;*.fb2;*.xps;*.oxps;*.cbz;*.txt)|*.pdf;*.epub;*.mobi;*.fb2;*.xps;*.oxps;*.cbz;*.txt|Todos los archivos (*.*)|*.*"
        with wx.FileDialog(self, "Seleccionar Archivo", wildcard=wildcard, style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST) as dlg:
            if dlg.ShowModal() == wx.ID_OK:
                self.txt_path.SetValue(dlg.GetPath())

    def OnRun(self, event):
        path = self.txt_path.GetValue()
        if not path: return
        self.btn_run.Disable()
        self.log_ctrl.Clear()
        threading.Thread(target=self.ProcessWorker, args=(path,), daemon=True).start()

    def AddTableToDoc(self, doc, matrix):
        if not matrix: return
        try:
            table = doc.add_table(rows=len(matrix), cols=len(matrix[0]))
            table.style = 'Table Grid'
            for i, row_data in enumerate(matrix):
                row_cells = table.rows[i].cells
                for j, cell_text in enumerate(row_data):
                    if j < len(row_cells):
                        row_cells[j].text = str(cell_text)
            doc.add_paragraph("") 
        except Exception as e:
            self.Log(f"Error insertando tabla en Word: {e}")

    def ProcessWorker(self, path):
        import concurrent.futures
        try:
            mode_idx = self.cb_mode.GetSelection() # 0: Transcribe, 1: Translate, 2: Hybrid
            
            eff_config = self.config.get("global", {}).copy()
            target_lang = eff_config.get("translate_to", "es")
            eff_config["translate_enabled"] = True
            eff_config["translate_type"] = "online" if eff_config.get("translate_type", "disabled") == "disabled" else eff_config.get("translate_type", "online")

            doc = fitz.open(path)
            total = len(doc)
            self.Log(f"Archivo abierto: {os.path.basename(path)}")
            self.Log(f"Total de páginas: {total}")
            
            output_type = "txt"
            if self.rb_word.GetValue(): output_type = "docx"
            elif self.rb_excel.GetValue(): output_type = "xlsx"
            
            output_path = os.path.join(os.path.dirname(path), os.path.splitext(os.path.basename(path))[0] + f"_procesado.{output_type}")
            
            word_doc = None
            if output_type == "docx":
                word_doc = Document()
            
            extracted_pages = []
            carry_over = ""

            # --- FASE 1: EXTRACCIÓN ---
            self.Log("--- FASE 1: EXTRACCIÓN ---")
            for i in range(total):
                wx.CallAfter(self.gauge.SetValue, int(((i+1)/total)*100))
                page = doc.load_page(i)
                
                page_data = {"page_num": i+1, "blocks": []} # blocks: [{'type': 'text'/'table', 'content': ...}]
                
                native_text = page.get_text("text").strip()
                
                if native_text and len(native_text) > 20:
                    self.Log(f"Página {i+1}: Texto nativo detectado.")
                    page_data["blocks"].append({"type": "text", "content": native_text})
                else:
                    self.Log(f"Página {i+1}: Imagen detectada. OCR...")
                    if not self.ocr:
                        self.ocr = OCREngine(self.config)
                        self.ocr.initialize()
                        
                    if self.chk_tables.GetValue() and not self.layout:
                        self.layout = LayoutEngine(self.config)
                        self.layout.initialize()
                        self.table = TableEngine(self.config)
                        self.table.initialize()
                        
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
                    img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR if pix.n == 3 else cv2.COLOR_RGBA2BGR)
                    
                    if self.chk_tables.GetValue() and self.layout:
                        regions = self.layout.analyze_page(img)
                        for reg in regions:
                            x1, y1, x2, y2 = [int(v) for v in reg['bbox']]
                            x1, y1 = max(0, x1), max(0, y1)
                            x2, y2 = min(img.shape[1], x2), min(img.shape[0], y2)
                            region_img = img[y1:y2, x1:x2]
                            if region_img.size == 0: continue
                            
                            if reg['type'] == 'table':
                                html = self.table.process_table_region(region_img)
                                matrix = self.table.html_to_matrix(html)
                                if matrix:
                                    page_data["blocks"].append({"type": "table", "content": matrix})
                            else:
                                elements = self.ocr.scan_image(region_img)
                                text_block = "\n".join([e.text for e in elements])
                                if text_block.strip():
                                    page_data["blocks"].append({"type": "text", "content": text_block})
                    else:
                        elements = self.ocr.scan_image(img)
                        text_page = "\n".join([e.text for e in elements])
                        if text_page.strip():
                            page_data["blocks"].append({"type": "text", "content": text_page})
                            
                # Procesar carry over de la página anterior
                if carry_over:
                    first_text = next((b for b in page_data["blocks"] if b["type"] == "text"), None)
                    if first_text and first_text["content"]:
                        join_char = " " if first_text["content"][0].islower() or first_text["content"][0] in ",;:" else "\n\n"
                        first_text["content"] = carry_over + join_char + first_text["content"]
                    else:
                        page_data["blocks"].insert(0, {"type": "text", "content": carry_over})
                    carry_over = ""

                # Evaluar si la página actual quedó cortada (solo miramos el último bloque si es de texto)
                if page_data["blocks"]:
                    last_block = page_data["blocks"][-1]
                    if last_block["type"] == "text":
                        content = last_block["content"].strip()
                        # Si no termina en un signo de puntuación de cierre
                        if content and not content[-1] in ".!?:\"'”)]>":
                            import re
                            matches = list(re.finditer(r'[.!?:]+(?=\s|$|\n)', content))
                            if matches:
                                split_idx = matches[-1].end()
                                carry_over = content[split_idx:].strip()
                                last_block["content"] = content[:split_idx].strip()
                            else:
                                carry_over = content
                                page_data["blocks"].pop()

                extracted_pages.append(page_data)

            # Si sobró algo en la última página, lo devolvemos
            if carry_over and extracted_pages:
                extracted_pages[-1]["blocks"].append({"type": "text", "content": carry_over})

            # --- FASE 2: TRADUCCIÓN ---
            if mode_idx in (1, 2):
                self.Log("--- FASE 2: TRADUCCIÓN ---")
                wx.CallAfter(self.gauge.SetValue, 0)
                
                t_type = eff_config.get("translate_type", "online")
                t_from = eff_config.get("translate_from", "en")
                t_to = eff_config.get("translate_to", "es")
                service = eff_config.get("translate_service", "google")
                swap = eff_config.get("translate_swap", False)
                parallel_pages = int(eff_config.get("translate_parallel_pages", 5))
                
                def translate_page(p_data):
                    # Unir todos los bloques de texto de la página
                    text_blocks = [b for b in p_data["blocks"] if b["type"] == "text"]
                    if text_blocks:
                        combined_text = "\n\n===SPLIT===\n\n".join([b["content"] for b in text_blocks])
                        translated_text = translator_instance.translate(combined_text, t_from, t_to, translate_type=t_type, service=service, swap=swap)
                        translated_parts = translated_text.split("\n\n===SPLIT===\n\n")
                        
                        # Asignar de vuelta (fallback si falla el split)
                        if len(translated_parts) == len(text_blocks):
                            for idx, b in enumerate(text_blocks):
                                b["translated_content"] = translated_parts[idx]
                        else:
                            text_blocks[0]["translated_content"] = translated_text
                            for b in text_blocks[1:]: b["translated_content"] = ""
                    
                    # Traducir tablas (por celda)
                    for b in p_data["blocks"]:
                        if b["type"] == "table":
                            t_matrix = []
                            for row in b["content"]:
                                t_row = []
                                for cell in row:
                                    if cell.strip():
                                        t_cell = translator_instance.translate(cell, t_from, t_to, translate_type=t_type, service=service, swap=swap)
                                        t_row.append(t_cell)
                                    else:
                                        t_row.append("")
                                t_matrix.append(t_row)
                            b["translated_content"] = t_matrix
                            
                    return p_data

                if t_type == "online" and parallel_pages > 1:
                    self.Log(f"Traduciendo con {parallel_pages} páginas en paralelo...")
                    with concurrent.futures.ThreadPoolExecutor(max_workers=parallel_pages) as executor:
                        futures = [executor.submit(translate_page, p) for p in extracted_pages]
                        for i, future in enumerate(futures):
                            extracted_pages[i] = future.result()
                            wx.CallAfter(self.gauge.SetValue, int(((i+1)/total)*100))
                else:
                    self.Log("Traduciendo páginas secuencialmente (Aceleración Lotes Offline activa)...")
                    for i, p in enumerate(extracted_pages):
                        extracted_pages[i] = translate_page(p)
                        wx.CallAfter(self.gauge.SetValue, int(((i+1)/total)*100))

            # --- FASE 3: ARMADO DEL ARCHIVO ---
            self.Log("--- FASE 3: ARMADO DE ARCHIVO ---")
            all_text_data = []
            
            for p_data in extracted_pages:
                for b in p_data["blocks"]:
                    original = b["content"]
                    translated = b.get("translated_content", "")
                    
                    if b["type"] == "text":
                        if mode_idx == 0:
                            final_text = original
                        elif mode_idx == 1:
                            final_text = translated
                        elif mode_idx == 2:
                            final_text = f"{original}\n{translated}\n"
                            
                        if output_type == "docx":
                            if mode_idx == 2:
                                word_doc.add_paragraph(original)
                                pg_trans = word_doc.add_paragraph(translated)
                                # italic formating if needed
                                word_doc.add_paragraph("")
                            else:
                                word_doc.add_paragraph(final_text)
                        elif output_type == "xlsx":
                            all_text_data.append({"Pagina": p_data["page_num"], "Tipo": "Texto", "Contenido": final_text})
                        else:
                            all_text_data.append(final_text + "\n\n")
                            
                    elif b["type"] == "table":
                        t_matrix = translated if mode_idx == 1 else original
                        # Híbrido en tablas: original + (traducido)
                        if mode_idx == 2 and translated:
                            t_matrix = []
                            for r_idx, row in enumerate(original):
                                new_row = []
                                for c_idx, cell in enumerate(row):
                                    t_c = translated[r_idx][c_idx] if r_idx < len(translated) and c_idx < len(translated[r_idx]) else ""
                                    new_row.append(f"{cell}\n({t_c})")
                                t_matrix.append(new_row)
                                
                        if output_type == "docx":
                            self.AddTableToDoc(word_doc, t_matrix)
                        elif output_type == "xlsx":
                            for row in t_matrix:
                                all_text_data.append({"Pagina": p_data["page_num"], "Tipo": "Tabla", "Contenido": " | ".join(row)})
                        else:
                            table_str = "\n".join([" | ".join(row) for row in t_matrix])
                            all_text_data.append(f"[TABLA PÁG {p_data['page_num']}]\n{table_str}\n")
                            
                if output_type == "docx" and p_data["page_num"] < total and self.chk_page_breaks.GetValue():
                    word_doc.add_page_break()

            self.Log("Guardando resultados...")
            if output_type == "docx":
                word_doc.save(output_path)
            elif output_type == "xlsx":
                pd.DataFrame(all_text_data).to_excel(output_path, index=False)
            else:
                with open(output_path, "w", encoding="utf-8") as f:
                    f.writelines(all_text_data)

            self.Log(f"¡ÉXITO! Archivo guardado en: {output_path}")
            wx.CallAfter(wx.MessageBox, "El documento ha sido procesado con éxito.", "Finalizado", wx.OK | wx.ICON_INFORMATION)
            
        except Exception as e:
            self.Log(f"ERROR CRÍTICO: {e}")
            import traceback
            traceback.print_exc()
            wx.CallAfter(wx.MessageBox, f"Ocurrió un error: {e}", "Error", wx.OK | wx.ICON_ERROR)
        finally:
            wx.CallAfter(self.btn_run.Enable)
            wx.CallAfter(self.gauge.SetValue, 0)

if __name__ == "__main__":
    app = wx.App()
    frame = TranscriptorFrame()
    frame.Show()
    app.MainLoop()
